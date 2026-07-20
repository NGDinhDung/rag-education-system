from pathlib import Path

import pymupdf
from docx import Document as DocxDocument


SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
}


def load_pdf(file_path: Path) -> str:
    text_parts: list[str] = []

    try:
        with pymupdf.open(file_path) as pdf:
            for page_number, page in enumerate(pdf, start=1):
                page_text = page.get_text(
                    "text",
                    sort=True,
                ).strip()

                if page_text:
                    text_parts.append(
                        f"[Trang {page_number}]\n{page_text}"
                    )

    except Exception as error:
        raise ValueError(
            f"Không thể đọc file PDF: {error}"
        ) from error

    return "\n\n".join(text_parts)


def load_docx(file_path: Path) -> str:
    text_parts: list[str] = []

    try:
        document = DocxDocument(file_path)

        for paragraph in document.paragraphs:
            paragraph_text = paragraph.text.strip()

            if paragraph_text:
                text_parts.append(paragraph_text)

        # Đọc thêm nội dung trong bảng DOCX
        for table in document.tables:
            for row in table.rows:
                row_values = [
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                ]

                if row_values:
                    text_parts.append(" | ".join(row_values))

    except Exception as error:
        raise ValueError(
            f"Không thể đọc file DOCX: {error}"
        ) from error

    return "\n\n".join(text_parts)


def load_txt(file_path: Path) -> str:
    encodings = (
        "utf-8",
        "utf-8-sig",
        "cp1258",
    )

    for encoding in encodings:
        try:
            return file_path.read_text(
                encoding=encoding
            ).strip()

        except UnicodeDecodeError:
            continue

        except Exception as error:
            raise ValueError(
                f"Không thể đọc file TXT: {error}"
            ) from error

    raise ValueError(
        "Không thể xác định bảng mã của file TXT."
    )


def load_document(file_path: str) -> str:
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(
            f"Không tìm thấy file: {file_path}"
        )

    extension = path.suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            "Chỉ hỗ trợ file PDF, DOCX và TXT."
        )

    if extension == ".pdf":
        content = load_pdf(path)

    elif extension == ".docx":
        content = load_docx(path)

    else:
        content = load_txt(path)

    if not content.strip():
        raise ValueError(
            "Không trích xuất được nội dung từ tài liệu."
        )

    return content